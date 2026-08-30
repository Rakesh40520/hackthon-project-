"""DOCX extraction."""
from __future__ import annotations

from typing import List

from app.document_processing.extractor import ExtractedDocument, ExtractedSection


def extract_docx(file_path: str, filename: str) -> ExtractedDocument:
    from docx import Document

    document = Document(file_path)
    sections: List[ExtractedSection] = []
    full_text_parts: List[str] = []
    current_section = "Body"
    for para in document.paragraphs:
        style = (para.style.name or "").lower() if para.style else ""
        text = para.text
        if not text.strip():
            continue
        if "heading" in style:
            current_section = text.strip()
        sections.append(ExtractedSection(text=text, section=current_section, document=filename))
        full_text_parts.append(text)

    tables: List[List[List[str]]] = []
    for t in document.tables:
        grid: List[List[str]] = []
        for row in t.rows:
            grid.append([cell.text.strip() for cell in row.cells])
        tables.append(grid)

    return ExtractedDocument(
        text="\n".join(full_text_parts),
        sections=sections,
        tables=tables,
        metadata={"filename": filename, "type": "docx"},
    )
