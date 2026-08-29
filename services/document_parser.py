"""Document text extraction for ClauseGuard.

Handles PDF (pypdf), DOCX (python-docx) and plain text. The extracted text is
the immutable "source of truth" that every AI claim gets verified against.
"""

from __future__ import annotations

import io
import os
from dataclasses import dataclass, field
from typing import List, Optional

SUPPORTED_EXTENSIONS = {".pdf": "PDF", ".txt": "TXT", ".md": "TXT", ".docx": "DOCX"}

# Below this many extracted characters we consider the parse useless.
MIN_USEFUL_CHARS = 120


@dataclass
class ParsedDocument:
    file_name: str = ""
    file_type: str = ""
    text: str = ""
    char_count: int = 0
    page_count: Optional[int] = None
    warnings: List[str] = field(default_factory=list)
    error: Optional[str] = None

    @property
    def ok(self) -> bool:
        return self.error is None and bool(self.text.strip())


def _parse_pdf(data: bytes, doc: ParsedDocument) -> None:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:
            doc.error = "The PDF is password-protected and cannot be read."
            return

    pages = len(reader.pages)
    doc.page_count = pages
    chunks: List[str] = []
    for index, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
            doc.warnings.append(f"Text could not be extracted from page {index + 1}.")
        chunks.append(page_text)
    doc.text = "\n\n".join(chunks)


def _parse_docx(data: bytes, doc: ParsedDocument) -> None:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts: List[str] = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    doc.text = "\n".join(parts)


def _parse_text(data: bytes, doc: ParsedDocument) -> None:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            doc.text = data.decode(encoding)
            if encoding != "utf-8-sig" and encoding != "utf-8":
                doc.warnings.append(f"File was decoded using {encoding}.")
            return
        except (UnicodeDecodeError, UnicodeError):
            continue
    doc.error = "Could not decode the text file with any common encoding."


def parse_document_bytes(file_name: str, data: bytes) -> ParsedDocument:
    """Parse raw file bytes into a ParsedDocument. Never raises."""
    ext = os.path.splitext(file_name or "")[1].lower()
    file_type = SUPPORTED_EXTENSIONS.get(ext)

    if file_type is None:
        return ParsedDocument(
            file_name=file_name or "",
            file_type="Unknown",
            error=(
                f"Unsupported file type '{ext or '(none)'}'. "
                "ClauseGuard supports PDF, TXT/MD and DOCX files."
            ),
        )

    doc = ParsedDocument(file_name=file_name, file_type=file_type)

    if not data:
        doc.error = "The uploaded file is empty."
        return doc

    try:
        if ext == ".pdf":
            _parse_pdf(data, doc)
        elif ext == ".docx":
            _parse_docx(data, doc)
        else:
            _parse_text(data, doc)
    except ModuleNotFoundError as exc:  # missing optional dependency
        doc.error = f"A required parsing library is missing: {exc}. Run pip install -r requirements.txt."
        return doc
    except Exception as exc:  # defensive: never crash the UI with a stack trace
        doc.error = f"Failed to extract text from the document: {exc}"
        return doc

    if doc.error:
        return doc

    doc.char_count = len(doc.text)

    if not doc.text.strip():
        if file_type == "PDF":
            doc.error = (
                "No extractable text was found in this PDF. It may be a scanned "
                "image; ClauseGuard requires text-based documents (OCR is not supported)."
            )
        else:
            doc.error = "The document contains no extractable text."
    elif doc.char_count < MIN_USEFUL_CHARS:
        doc.error = (
            f"Only {doc.char_count} characters of text were extracted — too little to "
            "analyze. The file may be corrupted or a scan."
        )

    return doc
